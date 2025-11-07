from classes.proposicao import Proposicao
from dataclasses import dataclass, asdict
from datetime import datetime
from libs import datas
from pathlib import Path
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from proposicoes_bd.msaccess import conexao_msaccess


def cria_estilo(lista_estilos, nome_estilo, alinhamento, fonte, tamanho, negrito=False, recuo_primeira_linha=None, recuo_esquerda=None):
    if nome_estilo in [s.name for s in lista_estilos]:
        return lista_estilos[nome_estilo]

    estilo = lista_estilos.add_style(nome_estilo, WD_STYLE_TYPE.PARAGRAPH)
    estilo.paragraph_format.alignment = alinhamento
    estilo.font.name = fonte
    estilo.font.size = tamanho
    estilo.font.bold = negrito
    if recuo_primeira_linha:
        estilo.paragraph_format.first_line_indent = recuo_primeira_linha
    if recuo_esquerda:
        estilo.paragraph_format.left_indent = recuo_esquerda
    return estilo


@dataclass
class Edital(Proposicao):
    lista_proposicoes: list = None
    usar_link: bool = True

    def gera_documento(self, arquivo_modelo, diretorio_geracao, banco_dados_proposicoes):
        try:
            documento = Document(arquivo_modelo)
            estilos = documento.styles

            stl_norm_just_pl16 = cria_estilo(
                estilos,
                'estilo_1',
                WD_ALIGN_PARAGRAPH.JUSTIFY,
                'Arial',
                Pt(12),
                negrito=False,
                recuo_primeira_linha=Cm(1.6)
            )

            db = conexao_msaccess(banco_dados_proposicoes)

            relator_anterior = ""
            for proposicao in self.lista_proposicoes:
                if relator_anterior != proposicao.relator:
                    documento.add_paragraph('', style=stl_norm_just_pl16)
                    documento.add_paragraph(f"Relator: {proposicao.relator}", style=stl_norm_just_pl16)
                    documento.add_paragraph('', style=stl_norm_just_pl16)

                numero_formatado = f"{proposicao.numero}/{proposicao.ano}"

                if self.usar_link:
                    registro_db = db.seleciona_formatado( numero_formatado)

                    if registro_db == None:
                        link = "www.alerj.rj.gov.br"
                    else:
                        link = registro_db[7]  #COLUNA LINK

                    p = documento.add_paragraph(f"{proposicao.ordem}) ", style=stl_norm_just_pl16)
                    texto_link = (
                        f"{'Emendas de Plenário ao Projeto de Lei nº ' if proposicao.emenda_de_plenario else 'Projeto de Lei nº '}" + numero_formatado
                    )
                    self.add_hyperlink(p, texto_link, link)

                    p.add_run(
                        f", do(s) Deputado(s) {proposicao.autores.title()}, que "
                        f"{proposicao.ementa}{'' if proposicao.ementa.endswith('.') else '.'}"
                    )
                else:
                    texto_link = (f"{'Emendas de Plenário ao Projeto de Lei nº ' if proposicao.emenda_de_plenario else 'Projeto de Lei nº '}" + numero_formatado)
                    p = documento.add_paragraph(f"{proposicao.ordem}) {texto_link}", style=stl_norm_just_pl16)
                    p.add_run(
                        f", do(s) Deputado(s) {proposicao.autores.title()}, que "
                        f"{proposicao.ementa}{'' if proposicao.ementa.endswith('.') else '.'}"
                    )

                documento.add_paragraph('', style=stl_norm_just_pl16)
                relator_anterior = proposicao.relator

            diretorio = Path(diretorio_geracao)
            diretorio.mkdir(parents=True, exist_ok=True)

            data_formatada = datetime.today().strftime('%Y%m%d_%H%M%S')
            nome_arquivo = f"edital_{data_formatada}.docx"
            caminho_arquivo = diretorio / nome_arquivo
            documento.save(caminho_arquivo)

            print(f"✅ Arquivo gerado: {caminho_arquivo}")

        except Exception as e:
            raise Exception(f"Verifique se a pasta ou o arquivo de modelo de Edital existe!\nDetalhes: {e}")


@dataclass
class Conclusao(Proposicao):
    arquivo_modelo: str = ""
    arquivo_modelo_voto_separado: str = ""
    diretorio_geracao: str = ""

    def gera_documento(self, data_sessao, reuniao):
        try:

            if self.parecer_vista and self.relator_vista:
                documento = Document(self.arquivo_modelo_voto_separado)
            else:
                documento = Document(self.arquivo_modelo)

            estilos = documento.styles

            stl_norm_just_pl16 = cria_estilo(
                estilos, 'estilo_1', WD_ALIGN_PARAGRAPH.JUSTIFY,
                'Arial', Pt(12), negrito=False, recuo_primeira_linha=Cm(1.6)
            )

            for paragrafo in documento.paragraphs:
                paragrafo.text = (
                    paragrafo.text
                    .replace('{{ NUMERO_PROJETO }}', f"{self.numero}/{self.ano}")
                    .replace('{{ REUNIAO }}', reuniao)
                    .replace('{{ DATA_REUNIAO }}', data_sessao)
                    .replace('{{ PARECER }}', self.parecer or "")
                    .replace('{{RELATOR_VOTO_SEPARADO}}', self.relator_vista or "")
                    .replace('{{PARECER_VOTO_SEPARADO}}', self.parecer_vista or "")
                )

                if '{{ TIPO_PROPOSICAO }}' in paragrafo.text:
                    tipo = (
                        f"à(s) Emenda(s) de Plenário ao {self.classifica_tipo_proposicao()}"
                        if self.emenda_de_plenario
                        else f"ao {self.classifica_tipo_proposicao()}"
                    )
                    paragrafo.text = paragrafo.text.replace('{{ TIPO_PROPOSICAO }}', tipo)

                if '{{ DATA_REUNIAO_POR_EXTENSO }}' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace(
                        '{{ DATA_REUNIAO_POR_EXTENSO }}',
                        datas.data_por_extenso(data_sessao)
                    )

                paragrafo.style = stl_norm_just_pl16

            diretorio = Path(self.diretorio_geracao)
            diretorio.mkdir(parents=True, exist_ok=True)

            eh_emenda = "EP " if self.emenda_de_plenario else ""
            nome_arquivo = (
                f"Conclusao {eh_emenda}{self.classifica_tipo_proposicao(True)}_"
                f"{self.numero}-{self.ano}.docx"
            )
            caminho_arquivo = diretorio / nome_arquivo
            documento.save(caminho_arquivo)

            print(f"✅ Arquivo gerado: {caminho_arquivo}")

        except Exception as e:
            raise Exception(f"Verifique se a pasta ou o arquivo de modelo de Conclusão existe!\nDetalhes: {e}")


def proposicao_para_edital(p: Proposicao) -> Edital:
    return Edital(**asdict(p))


def proposicao_para_conclusao(p: Proposicao) -> Conclusao:
    return Conclusao(**asdict(p))
