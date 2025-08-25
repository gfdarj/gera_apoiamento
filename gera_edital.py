from classes.aplicacao import Configuracao
from classes.planilha import PlanilhaProjetos
from classes.documento import cria_estilo
import locale
from docx import Document
from datetime import datetime
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm


# Carrega as proposições
print("Teste de execução")
P = PlanilhaProjetos()
print("Carregando projetos")
proposicoes = P.CarregaColunas()
print(f"Projetos selecionados: {len(proposicoes)}")


# Ordena as proposições
if len(proposicoes) > 0:
    proposicoes = sorted(proposicoes, key=lambda x: (x.relator, int(x.ano), int(x.numero)))

for d in proposicoes:
    print(f"relator: {d.relator}  ------ numero:{d.numero}/{d.ano}  ----- EP? {d.emenda_de_plenario}")


config = Configuracao()


#print(config.arquivo_modelo_edital)
documento = Document(config.arquivo_modelo_edital)
estilos = documento.styles

# Texto em geral e linhas em branco
# Fonte Arial 12, normal
# Parágrafo justificado, primeira linha com recuo de 1.6 cm
stl_norm_just_pl16 = cria_estilo(estilos, 'estilo_1', WD_ALIGN_PARAGRAPH.JUSTIFY, 'Arial', Pt(12), negrito=False, recuo_primeira_linha=Cm(1.6), recuo_esquerda=None)


linha = 1
relator = ""
for proposicao in proposicoes:

    if relator == "" or relator != proposicao.relator:
        documento.add_paragraph('', style=stl_norm_just_pl16)
        documento.add_paragraph(f"Relator: {proposicao.relator}", style=stl_norm_just_pl16)
        documento.add_paragraph('', style=stl_norm_just_pl16)

    texto = f"{linha}) {'Emendas de Plenário ao Projeto de Lei nº ' if proposicao.emenda_de_plenario else 'Projeto de Lei nº '}{proposicao.numero}/{proposicao.ano}"
    texto += f", do(s) Deputado(s) {proposicao.autores.title()}, que {proposicao.ementa}{"" if proposicao.ementa[-1] == "." else "."}"
    documento.add_paragraph(texto, style=stl_norm_just_pl16)
    documento.add_paragraph('', style=stl_norm_just_pl16)
    linha += 1
    relator = proposicao.relator


nome_arquivo = 'edital_proximo.docx'
documento.save(nome_arquivo)





